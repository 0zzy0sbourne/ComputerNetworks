from docx import Document

def extract_pages(input_file, output_file, start_page, end_page):
    doc = Document(input_file)

    if start_page < 1 or end_page > len(doc.paragraphs):
        print("Invalid page range")
        return

    new_doc = Document()

    for page in range(start_page - 1, end_page):
        new_doc.add_paragraph(doc.paragraphs[page].text)

    new_doc.save(output_file)
    print(f"Pages {start_page} to {end_page} extracted and saved to {output_file}")

if __name__ == "__main__":
    input_file = "Solutions-7th-Edition.docx"  # Replace with your input .docx file
    output_file = "solutions_extracted.docx"  # Replace with the desired output file
    start_page = 1  # Replace with the starting page number
    end_page = 70  # Replace with the ending page number

    extract_pages(input_file, output_file, start_page, end_page)
